import os
import numpy as np
from langchain.tools import Tool
from langchain.agents import initialize_agent, AgentType
from langchain_community.utilities import SQLDatabase
from langchain_community.agent_toolkits.sql.base import create_sql_agent
from langchain.chains import RetrievalQA
from langchain.schema import Document, BaseRetriever
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS as LangchainFAISS
from docx import Document as DocxDocument
from openpyxl import Workbook
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from langchain.schema import SystemMessage
import json
import re
import time
import hashlib
import streamlit as st
from dotenv import load_dotenv
import requests
import uuid
from langchain.memory import ConversationBufferMemory
from chat_history import load_all_conversations
from prompt import SYSTEM_PROMPT
import torch
import socket
from typing import Dict, Any, List
import pathlib
import threading


secret_path = "/etc/secrets/OPENAI_API_KEY"
if os.path.exists(secret_path):
    with open(secret_path) as f:
        os.environ["OPENAI_API_KEY"] = f.read().strip()


if not os.getenv("OPENAI_API_KEY"):
    st.error("OPENAI_API_KEY is missing at runtime.")
    st.stop()
# Load environment variables from .env file
# load_dotenv()

# Access MySQL credentials
# MYSQL_HOST = os.getenv("MYSQL_HOST")
# MYSQL_USER = os.getenv("MYSQL_USER")
# MYSQL_PASSWORD = os.getenv("MYSQL_PASSWORD")
# MYSQL_DATABASE = os.getenv("MYSQL_DATABASE")
# MYSQL_PORT = os.getenv("MYSQL_PORT", "3306")

# MYSQL_HOST='veggiesmart.sct-lb.net'
# MYSQL_USER='hxaveggies_user'
# MYSQL_PASSWORD='hxaveggies_user_pass'
# MYSQL_DATABASE='hxaveggies_db'
# MYSQL_PORT=3306
# MYSQL_HOST = st.secrets["mysql"]["MYSQL_HOST"]
# MYSQL_USER = st.secrets["mysql"]["MYSQL_USER"]
# MYSQL_PASSWORD = st.secrets["mysql"]["MYSQL_PASSWORD"]
# MYSQL_DATABASE = st.secrets["mysql"]["MYSQL_DATABASE"]
# MYSQL_PORT = st.secrets["mysql"]["MYSQL_PORT"]
MYSQL_HOST = os.environ.get("MYSQL_HOST")
MYSQL_PORT = int(os.environ.get("MYSQL_PORT", 3306))
MYSQL_USER = os.environ.get("MYSQL_USER")
MYSQL_PASSWORD = os.environ.get("MYSQL_PASSWORD")
MYSQL_DATABASE = os.environ.get("MYSQL_DATABASE")

# st.write({
#     "MYSQL_HOST": MYSQL_HOST,
#     "MYSQL_PORT": MYSQL_PORT,
#     "MYSQL_USER": MYSQL_USER,
#     "MYSQL_DATABASE": MYSQL_DATABASE,
#     "MYSQL_PASSWORD": "SET" if MYSQL_PASSWORD else None
# })

# try:
#     socket.gethostbyname(MYSQL_HOST)
#     st.success("DNS resolved OK")
# except Exception as e:
#     st.error(f"DNS resolution failed: {e}")

# try:
#     conn = pymysql.connect(
#         host=MYSQL_HOST,
#         user=MYSQL_USER,
#         password=MYSQL_PASSWORD,
#         database=MYSQL_DATABASE,
#         port=int(MYSQL_PORT),
#         connect_timeout=10
#     )
#     st.success("Raw PyMySQL connection successful")
#     conn.close()
# except Exception as e:
#     st.error(f"PyMySQL connection failed: {e}")
# Access API credentials
# API_BASE_URL = os.getenv("API_BASE_URL", "http://192.168.10.82/hxa/ai_api/index.php")
# API_KEY = os.getenv("API_KEY")
# db_user = st.secrets["mysql"]["MYSQL_USER"] 
# db_password = st.secrets["mysql"]["MYSQL_PASSWORD"] 
# db_host = st.secrets["mysql"]["MYSQL_HOST"] 
# db_port = st.secrets["mysql"]["MYSQL_PORT"] 
# db_name = st.secrets["mysql"]["MYSQL_DATABASE"]
# Access OpenAI API key
# api_key = os.getenv("OPENAI_API_KEY")
# os.environ["OPENAI_API_KEY"] = api_key

# Initialize SentenceTransformer
# embedder = SentenceTransformer('all-MiniLM-L6-v2', device='cpu')

# heavy imports moved into functions to reduce cold-start time


@st.cache_resource
def get_llm(model: str = "gpt-4.1-nano", temperature: float = 0):
    # lazy import ChatOpenAI to avoid heavy import cost at module load
    from langchain_openai import ChatOpenAI
    return ChatOpenAI(model=model, temperature=temperature)


llm = get_llm()

# ---------- Index & embedding cache settings ----------
INDEX_DIR = os.path.join("data", ".index")
INDEX_FAISS_PATH = os.path.join(INDEX_DIR, "index.faiss")
INDEX_DOCS_PATH = os.path.join(INDEX_DIR, "docs.json")
INDEX_META_PATH = os.path.join(INDEX_DIR, "meta.json")
INDEX_HASHES_PATH = os.path.join(INDEX_DIR, "file_hashes.json")

# Singleton embedding wrapper (LangChain HuggingFaceEmbeddings)
_EMBEDDING_WRAPPER = None

def get_embedding_wrapper(model_name: str = 'sentence-transformers/all-MiniLM-L6-v2'):
    global _EMBEDDING_WRAPPER
    if _EMBEDDING_WRAPPER is None:
        _EMBEDDING_WRAPPER = get_cpu_huggingface_embeddings(model_name)
    return _EMBEDDING_WRAPPER

def _compute_file_hashes(data_folder: str = "data") -> Dict[str, Dict[str, int]]:
    """
    Return a fast signature for files in `data_folder` keyed by filename -> {mtime, size}.
    This avoids reading entire file contents and is much faster for large files.
    """
    sigs = {}
    try:
        for file in os.listdir(data_folder):
            file_path = os.path.join(data_folder, file)
            if os.path.isfile(file_path):
                try:
                    stt = os.stat(file_path)
                    sigs[file] = {"mtime": int(stt.st_mtime), "size": stt.st_size}
                except Exception:
                    continue
    except Exception:
        pass
    return sigs


def _safe_name(filename: str) -> str:
    return hashlib.sha1(filename.encode("utf-8")).hexdigest()


def _per_file_paths(filename: str):
    safe = _safe_name(filename)
    base = os.path.join(INDEX_DIR, safe)
    pathlib.Path(base).mkdir(parents=True, exist_ok=True)
    return {
        "dir": base,
        "docs": os.path.join(base, "docs.json"),
        "meta": os.path.join(base, "meta.json"),
        "emb": os.path.join(base, "emb.npy"),
        "hash": os.path.join(base, "hash.txt"),
    }


# Background-threaded indexing removed — rely on `@st.cache_resource` for cached index builds


# Cached global index wrapper using file-hash as cache key
@st.cache_resource
def cached_build_index(file_sig: str, data_folder="data", chunk_size: int = 300, overlap: int = 20, batch_size: int = 64):
    # file_sig should be a deterministic but cheap fingerprint (mtime+size) of files
    return build_index(data_folder=data_folder, chunk_size=chunk_size, overlap=overlap, batch_size=batch_size)



# -------------------------------
# SentenceTransformer helpers
# -------------------------------
def get_cpu_huggingface_embeddings(
    model_name: str = "sentence-transformers/all-MiniLM-L6-v2",
):
    """LangChain wrapper – always safe on CPU."""
    from langchain.embeddings import HuggingFaceEmbeddings

    return HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )


def get_cpu_sentence_transformer(model_name: str = "all-MiniLM-L6-v2"):
    """SentenceTransformer – never calls .to() on a meta tensor."""
    from sentence_transformers import SentenceTransformer
    import torch

    model = SentenceTransformer(model_name)               # load without device
    if any(p.device.type == "meta" for p in model.parameters()):
        model = SentenceTransformer(model_name, device="cpu")
    else:
        model = model.to(torch.device("cpu"))
    return model

def extract_text_from_pdf(pdf_path):
    # import here to avoid heavy import at module load
    import fitz
    text = ""
    with fitz.open(pdf_path) as doc:
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_structured_from_excel(excel_path):
    import pandas as pd
    df = pd.read_excel(excel_path)
    df = df.fillna("")
    structured_rows = []
    text_rows = []
    for _, row in df.iterrows():
        row_dict = {col: str(row[col]) for col in df.columns}
        row_text = " | ".join([f"{col}: {row_dict[col]}" for col in df.columns])
        structured_rows.append(row_dict)
        text_rows.append(row_text)
    return structured_rows, "\n".join(text_rows)

def chunk_text(text, chunk_size=300, overlap=50):
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        chunks.append(chunk)
        start += chunk_size - overlap
    return chunks

def build_index(data_folder="data", chunk_size: int = 300, overlap: int = 20, batch_size: int = 64):
    docs, metadata = [], []

    pathlib.Path(INDEX_DIR).mkdir(parents=True, exist_ok=True)
    # lazy import faiss to avoid import cost at module load
    import faiss

    # Quick-check: if a saved global index exists and files haven't changed, load it
    try:
        current_hashes = _compute_file_hashes(data_folder)
        if os.path.exists(INDEX_FAISS_PATH) and os.path.exists(INDEX_DOCS_PATH) and os.path.exists(INDEX_META_PATH) and os.path.exists(INDEX_HASHES_PATH):
            try:
                with open(INDEX_HASHES_PATH, "r", encoding="utf-8") as fh:
                    saved_hashes = json.load(fh)
                if saved_hashes == current_hashes:
                    # load persisted index & metadata
                    with open(INDEX_DOCS_PATH, "r", encoding="utf-8") as fh:
                        docs = json.load(fh)
                    with open(INDEX_META_PATH, "r", encoding="utf-8") as fh:
                        metadata = json.load(fh)
                    index = faiss.read_index(INDEX_FAISS_PATH)
                    return docs, metadata, index
            except Exception:
                # If any load step fails, fall back to rebuilding
                pass
    except Exception:
        # ignore hashing errors and rebuild
        pass

    # We'll build per-file embeddings (cached) and then combine
    embedder = get_embedding_wrapper()
    all_embs = []
    doc_list = []
    meta_list = []

    files = [f for f in os.listdir(data_folder) if os.path.isfile(os.path.join(data_folder, f))]
    total_files = len(files)
    file_idx = 0

    for file in files:
        file_idx += 1
        file_path = os.path.join(data_folder, file)
        paths = _per_file_paths(file)

        # compute a cheap file signature (mtime:size) and compare with cached signature
        try:
            stt = os.stat(file_path)
            file_sig = f"{int(stt.st_mtime)}:{stt.st_size}"
        except Exception:
            file_sig = None

        use_cache = False
        if file_sig and os.path.exists(paths["hash"]):
            try:
                with open(paths["hash"], "r", encoding="utf-8") as fh:
                    saved = fh.read().strip()
                if saved == file_sig and os.path.exists(paths["emb"]) and os.path.exists(paths["docs"]) and os.path.exists(paths["meta"]):
                    use_cache = True
            except Exception:
                use_cache = False

        if use_cache:
            try:
                file_docs = json.load(open(paths["docs"], "r", encoding="utf-8"))
                file_meta = json.load(open(paths["meta"], "r", encoding="utf-8"))
                file_emb = np.load(paths["emb"])
            except Exception:
                use_cache = False

        if not use_cache:
            # extract text and chunk depending on filetype
            if file.lower().endswith(".pdf"):
                text = extract_text_from_pdf(file_path)
                chunks = chunk_text(text, chunk_size=chunk_size, overlap=overlap)
                file_docs = [c.strip() for c in chunks if c.strip()]
                file_meta = [{"source": file, "type": "pdf", "structured": None} for _ in file_docs]
            elif file.lower().endswith((".xlsx", ".xls")):
                structured_rows, text_data = extract_structured_from_excel(file_path)
                chunks = chunk_text(text_data, chunk_size=chunk_size, overlap=overlap)
                file_docs = [c.strip() for c in chunks if c.strip()]
                file_meta = [{"source": file, "type": "excel", "structured": structured_rows} for _ in file_docs]
            elif file.lower().endswith((".txt", ".md")):
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        text = f.read()
                except Exception:
                    continue
                chunks = chunk_text(text, chunk_size=chunk_size, overlap=overlap)
                file_docs = [c.strip() for c in chunks if c.strip()]
                file_meta = [{"source": file, "type": "text", "structured": None} for _ in file_docs]
            else:
                # unsupported file types are skipped
                continue

            # compute embeddings for this file in batches
            emb_list = []
            for i in range(0, len(file_docs), batch_size):
                batch = file_docs[i : i + batch_size]
                try:
                    batch_emb = embedder.embed_documents(batch)
                except Exception:
                    batch_emb = []
                    for d in batch:
                        batch_emb.append(embedder.embed_documents([d])[0])
                emb_list.extend(batch_emb)

            file_emb = np.array(emb_list, dtype=np.float32)

            # persist per-file cache (store cheap file signature)
            try:
                np.save(paths["emb"], file_emb)
                with open(paths["docs"], "w", encoding="utf-8") as fh:
                    json.dump(file_docs, fh)
                with open(paths["meta"], "w", encoding="utf-8") as fh:
                    json.dump(file_meta, fh)
                if file_sig:
                    with open(paths["hash"], "w", encoding="utf-8") as fh:
                        fh.write(file_sig)
            except Exception:
                pass

        # append to global lists
        if len(file_docs) > 0:
            start_idx = len(doc_list)
            doc_list.extend(file_docs)
            meta_list.extend(file_meta)
            all_embs.append(file_emb)

        # update quick progress in session_state if present
        try:
            st.session_state["index_progress"] = file_idx / max(1, total_files)
        except Exception:
            pass

    if not doc_list:
        dim = 384
        index = faiss.IndexFlatIP(dim)
        return [], [], index

    # concatenate embeddings
    embeddings = np.vstack(all_embs)

    dim = embeddings.shape[1]
    index = faiss.IndexFlatIP(dim)
    index.add(embeddings)

    # Persist global index and metadata for fast reloads
    try:
        faiss.write_index(index, INDEX_FAISS_PATH)
        with open(INDEX_DOCS_PATH, "w", encoding="utf-8") as fh:
            json.dump(doc_list, fh)
        with open(INDEX_META_PATH, "w", encoding="utf-8") as fh:
            json.dump(meta_list, fh)
        with open(INDEX_HASHES_PATH, "w", encoding="utf-8") as fh:
            json.dump(_compute_file_hashes(data_folder), fh)
    except Exception:
        pass

    return doc_list, meta_list, index

# def get_api_tool(memory=None):
#     """
#     ERP API tool that dynamically fetches raw JSON data.
#     No counts, sums, or aggregations are computed in the code.
#     The LLM must interpret the raw JSON itself to compute counts, sums, maximums, or any other metrics.
#     """
#     def call_api(query: str) -> str:
#         try:
#             # Step 1: Ask LLM which endpoints to call
#             classification_prompt = f"""
#                 You are a strict JSON generator.
#                 Return ONLY a JSON list (no text, no explanation).
#                 Options: {list(API_ENDPOINTS.keys())}
#                 Query: {query}
#                 Example: ["invoices", "sales_orders"]
#                 """
#             classification = llm.invoke(classification_prompt).content.strip()
#             try:
#                 endpoints = json.loads(classification)
#             except json.JSONDecodeError:
#                 return "No relevant data found in the sources."

#             if not endpoints:
#                 return "No relevant data found in the sources."

#             # Step 2: Call all selected endpoints
#             results = {}
#             for endpoint in endpoints:
#                 if endpoint not in API_ENDPOINTS:
#                     results[endpoint] = {"error": f"Endpoint {endpoint} not found."}
#                     continue

#                 config = API_ENDPOINTS[endpoint]
#                 headers = {
#                     "Authorization": f"Bearer {API_KEY}" if API_KEY else "",
#                     "API-Key": API_KEY if API_KEY else "",
#                     "Content-Type": "application/json"
#                 }

#                 try:
#                     response = requests.get(config["url"], headers=headers, timeout=10)
#                     response.raise_for_status()
#                     data = response.json()
#                     results[endpoint] = data
#                 except requests.exceptions.RequestException as e:
#                     results[endpoint] = {"error": f"API unavailable for {endpoint}: {str(e)}"}

#             # Step 3: Hand raw data directly to LLM
#             llm_prompt =  f"""
#                 JSON DATA: {json.dumps(results)}

#                 Question: {query}

#                 Fill all numeric fields exactly from the JSON provided. 
#                 Do not skip any entries. Do not summarize. 
#                 explain the result.
#                 """
#             answer = llm.invoke(llm_prompt).content.strip()
#             return f"FINAL_ANSWER: {answer}"

#         except Exception as e:
#             return f"Unexpected error in API tool: {str(e)}"

#     return Tool(
#         name="erp_api_tool",
#         func=call_api,
#         description="Fetches raw ERP data (JSON). No pre-computed metrics — the AI must compute counts, sums, maximums, etc. dynamically."
#     )

def get_sqlite_tool(memory=None):
    db = SQLDatabase.from_uri("sqlite:///structured_data.db")
    system_message = SystemMessage(content=SYSTEM_PROMPT)
    sql_agent_executor = create_sql_agent(
        llm=llm,
        db=db,
        agent_type=AgentType.OPENAI_FUNCTIONS,
        memory=memory,
        verbose=True,
        agent_kwargs={"system_message": system_message}
    )

    return Tool(
        name="sqlite_tool",
        func=sql_agent_executor.run,
        description="Answer questions related to user's chat history stored in SQLite."
    )


def generate_table_info(engine):
    inspector = inspect(engine)
    table_info = {}

    for table_name in inspector.get_table_names():
        columns = inspector.get_columns(table_name)
        column_lines = [f"    - {col['name']} ({col['type']})" for col in columns]
        description = f"### {table_name}\n- Fields:\n" + "\n".join(column_lines)
        table_info[table_name] = description

    return table_info

def get_mysql_tool(memory=None, db_uri=None):
    if db_uri is None:
        db_uri = f"mysql+pymysql://{MYSQL_USER}:{MYSQL_PASSWORD}@{MYSQL_HOST}:{MYSQL_PORT}/{MYSQL_DATABASE}"
    # The tool function is lazy: it will create DB engine and SQL agent only when invoked.
    table_info = {
        "tbl_purchase_order": """
        ### tbl_purchase_order
        - This table contains all order data.
        - Sales Orders: is_sales = 1
        - Purchase Orders: is_sales = 0 and is_quotation = 0
        - Quotations: is_quotation = 1
        """,
        "tbl_invoices": """
        ### tbl_invoices
        - Contains invoice information.
        """,
        "tbl_invoice_products_details": """
        ### tbl_invoice_products_details
        - Contains invoice line items with Product_ID and Product_Quantity.
        - Must join with tbl_stock_products to get Product_Name for Product_ID.
        """,
        "tbl_stock_products": """
        ### tbl_stock_products
        - Contains product information with Product_ID and Product_Name.
        """,
    }

    def _call_mysql_tool(query: str) -> str:
        try:
            # lazy imports
            from sqlalchemy import create_engine, inspect
            from langchain_community.utilities import SQLDatabase
            from langchain_community.agent_toolkits.sql.base import create_sql_agent
            from langchain.schema import SystemMessage

            engine = create_engine(db_uri, connect_args={"connect_timeout": 5})
            inspector = inspect(engine)
            existing_tables = inspector.get_table_names()

            # Only keep tables that exist in both the DB and your table_info
            valid_tables = [t for t in table_info.keys() if t in existing_tables]

            db = SQLDatabase.from_uri(
                db_uri,
                include_tables=valid_tables,
                sample_rows_in_table_info=3,
                custom_table_info=table_info
            )

            system_message = SystemMessage(content=SYSTEM_PROMPT)
            sql_agent_executor = create_sql_agent(
                llm=llm,
                db=db,
                agent_type=AgentType.OPENAI_FUNCTIONS,
                memory=memory,
                verbose=True,
                agent_kwargs={"system_message": system_message}
            )

            return sql_agent_executor.run(query)
        except Exception as e:
            return f"MySQL tool error: {e}"

    return Tool(
        name="mysql_tool",
        func=_call_mysql_tool,
        description="Answer questions about structured financial data stored in MySQL (via phpMyAdmin). Tool creates DB connection lazily when invoked."
    )


def list_mysql_tables():
    db_uri = f"mysql+pymysql://{MYSQL_USER}:{MYSQL_PASSWORD}@{MYSQL_HOST}:{MYSQL_PORT}/{MYSQL_DATABASE}"
    from sqlalchemy import create_engine, inspect
    engine = create_engine(db_uri)
    inspector = inspect(engine)
    print("Tables in DB:", inspector.get_table_names())

def get_retriever_tool(docs, metadata, memory=None):
    # Use a local retriever that reuses a persisted FAISS index when available
    class LocalRetriever(BaseRetriever):
        index: Any
        docs: List[str]
        metadata: List[Dict]
        embedder: Any
        k: int = 4

        def get_relevant_documents(self, query):
            try:
                q_emb = self.embedder.embed_query(query)
            except Exception:
                q_emb = self.embedder.embed_documents([query])[0]

            q_emb = np.array(q_emb, dtype=np.float32)
            if q_emb.ndim == 1:
                q_emb = q_emb.reshape(1, -1)

            D, I = self.index.search(q_emb, self.k)
            results = []
            for idx in I[0]:
                if idx < len(self.docs):
                    results.append(Document(page_content=self.docs[idx], metadata=self.metadata[idx]))
            return results

    # Build or load index (fast when persisted). Use cached global index when possible.
    file_hashes = _compute_file_hashes("data")
    file_sig = json.dumps(file_hashes, sort_keys=True)

    docs_cached, meta_cached, index = cached_build_index(file_sig, data_folder="data")

    # If a session memory is provided, do NOT cache QA chain (avoid caching with memory)
    if memory is not None:
        import faiss
        embedding = get_embedding_wrapper()
        retr = LocalRetriever(index=index, docs=docs_cached, metadata=meta_cached, embedder=embedding)
        qa_chain = RetrievalQA.from_chain_type(llm=llm, retriever=retr, memory=memory, return_source_documents=False)
        return Tool(name="document_retriever", func=qa_chain.run, description="Document retriever (non-cached, per-session memory)")

    # memory is None -> safe to cache QA chain globally
    cached = st.session_state.get("qa_chain")
    if cached and getattr(cached, "_index_sig", None) == file_sig:
        return Tool(name="document_retriever", func=cached.run, description="Cached document retriever")

    @st.cache_resource
    def _build_qa_chain(index_hash: str):
        import faiss
        embedding = get_embedding_wrapper()
        retr = LocalRetriever(index=index, docs=docs_cached, metadata=meta_cached, embedder=embedding)
        qa = RetrievalQA.from_chain_type(llm=llm, retriever=retr, memory=None, return_source_documents=False)
        qa._index_sig = index_hash
        return qa

    qa_chain = _build_qa_chain(file_sig)
    st.session_state["qa_chain"] = qa_chain

    return Tool(
        name="document_retriever",
        func=qa_chain.run,
        description="Useful for answering questions from resumes, documents, or ERP policies. (cached)"
    )
def get_multi_agent(_, docs, metadata, db_uri=None, memory=None, conversation_history=None):
    if memory is None:
        memory = ConversationBufferMemory(return_messages=True, memory_key="chat_history")

    if conversation_history and not memory.chat_memory.messages:
        for msg in conversation_history:
            if msg["role"] == "user":
                memory.chat_memory.add_user_message(msg["content"])
            elif msg["role"] == "assistant":
                memory.chat_memory.add_ai_message(msg["content"])

    tools = [
        get_sqlite_tool(memory),
        get_retriever_tool(docs, metadata, memory),
    ]

    # Add MySQL tool lazily if host configured (no upfront connection test)
    if MYSQL_HOST:
        if db_uri is None:
            db_uri = (
                f"mysql+pymysql://{MYSQL_USER}:{MYSQL_PASSWORD}"
                f"@{MYSQL_HOST}:{MYSQL_PORT}/{MYSQL_DATABASE}"
            )
        tools.insert(0, get_mysql_tool(memory, db_uri))

    return initialize_agent(
        tools=tools,
        llm=llm,
        agent=AgentType.OPENAI_FUNCTIONS,
        memory=memory,
        verbose=True,
        handle_parsing_errors=True,
        agent_kwargs={"system_message": SystemMessage(content=SYSTEM_PROMPT)}
    )

def generate_chart(query: str, agent=None, retrieved_info: str = "") -> dict:
    default_response = {
        "chart_type": None,
        "chart_data": None,
        "render_method": "plotly",
        "error": "No data available for chart generation"
    }

    # Check if retrieved_info contains valid chart data
    try:
        chart_data = json.loads(retrieved_info) if retrieved_info else {}
        if chart_data.get("labels") and chart_data.get("values") and chart_data.get("title"):
            # Use provided chart data directly for Chart.js
            chartjs_config = {
                "type": "bar",  # Default to bar, adjust based on query if needed
                "data": {
                    "labels": chart_data["labels"],
                    "datasets": [{
                        "label": chart_data.get("title", "Data"),
                        "data": chart_data["values"],
                        "backgroundColor": [
                            "#83c5be", "#006d77", "#ff6b6b", "#4a5568", "#ffd60a",
                            "#8338ec", "#3a86ff", "#f72585"
                        ],
                        "borderColor": "#2d3748",
                        "borderWidth": 1
                    }]
                },
                "options": {
                    "responsive": True,
                    "plugins": {
                        "legend": {"position": "top"},
                        "title": {"display": True, "text": chart_data.get("title", "Chart")}
                    },
                    "scales": {
                        "x": {"title": {"display": True, "text": "Categories"}},
                        "y": {"title": {"display": True, "text": "Values"}}
                    }
                }
            }
            return {
                "chart_type": "bar",
                "chart_data": chartjs_config,
                "render_method": "chartjs",
                "error": None
            }
    except json.JSONDecodeError:
        pass  # Proceed to query the agent if retrieved_info is not valid JSON

    chart_types = {
        "bar": ["bar", "column", "histogram"],
        "line": ["line", "trend"],
        "pie": ["pie", "doughnut"],
        "scatter": ["scatter", "point"],
        "area": ["area"]
    }
    selected_chart_type = None
    for chart_type, keywords in chart_types.items():
        if any(keyword in query.lower() for keyword in keywords):
            selected_chart_type = chart_type
            break
    if not selected_chart_type:
        selected_chart_type = "bar"

    if agent:
        chart_prompt = f"""
        Query the available data sources (preferably get_mysql_tool, then SQL database or document index) to retrieve structured data for the query: '{query}'.
        For product-related queries, always join tbl_invoice_products_details with tbl_stock_products to include Product_Name.
        Return the data in a JSON format suitable for a {selected_chart_type} chart, with:
        - 'labels': list of strings for x-axis or categories
        - 'values': list of numbers for y-axis or data points
        - 'title': string for chart title
        If no relevant data is found, return an empty JSON: {{}}
        """
        try:
            response = agent.invoke({"input": chart_prompt}).get("output", "{}")
            chart_data = json.loads(response)
        except Exception as e:
            return {**default_response, "error": f"Error generating chart data: {e}"}
    else:
        chart_data = {}

    if not chart_data or not chart_data.get("labels") or not chart_data.get("values"):
        return default_response

    # lazy import plotly to avoid heavy import at module load
    import plotly.express as px

    if selected_chart_type == "bar":
        fig = px.bar(x=chart_data["labels"], y=chart_data["values"], title=chart_data.get("title", "Chart"))
    elif selected_chart_type == "line":
        fig = px.line(x=chart_data["labels"], y=chart_data["values"], title=chart_data.get("title", "Chart"))
    elif selected_chart_type == "pie":
        fig = px.pie(names=chart_data["labels"], values=chart_data["values"], title=chart_data.get("title", "Chart"))
    elif selected_chart_type == "scatter":
        fig = px.scatter(x=chart_data["labels"], y=chart_data["values"], title=chart_data.get("title", "Chart"))
    else:
        return {**default_response, "error": f"Unsupported chart type: {selected_chart_type}"}

    chartjs_config = {
        "type": selected_chart_type,
        "data": {
            "labels": chart_data["labels"],
            "datasets": [{
                "label": chart_data.get("title", "Data"),
                "data": chart_data["values"],
                "backgroundColor": [
                    "#83c5be", "#006d77", "#ff6b6b", "#4a5568", "#ffd60a",
                    "#8338ec", "#3a86ff", "#f72585"
                ],
                "borderColor": "#2d3748",
                "borderWidth": 1
            }]
        },
        "options": {
            "responsive": True,
            "plugins": {
                "legend": {"position": "top"},
                "title": {"display": True, "text": chart_data.get("title", "Chart")}
            },
            "scales": {
                "x": {"title": {"display": True, "text": "Categories"}},
                "y": {"title": {"display": True, "text": "Values"}}
            } if selected_chart_type in ["bar", "line", "scatter"] else {}
        }
    }

    return {
        "chart_type": selected_chart_type,
        "chart_data": chartjs_config if "chartjs" in query.lower() else fig,
        "render_method": "chartjs" if "chartjs" in query.lower() else "plotly",
        "error": None
    }
def generate_document(query: str, file_format: str = "pdf", agent=None, retrieved_info: str = "") -> str:
    timestamp = str(int(time.time()))
    query_hash = hashlib.md5(query.encode()).hexdigest()[:8]
    file_path = f"generated_report_{timestamp}_{query_hash}.{file_format.lower()}"

    if file_format.lower() in ["docx", "pdf"]:
        if agent:
            agent_prompt = f"""
            Generate concise, factual content for a {file_format.upper()} document based strictly on: '{query}'.
            Use the get_mysql_tool for real-time ERP data, falling back to other tools only if necessary.
            """
            content = agent.invoke({"input": agent_prompt}).get("output", "Insufficient information available")
        else:
            if not retrieved_info.strip():
                retrieved_info = "No relevant information found in the available data sources."
            content = retrieved_info

        if file_format.lower() == "docx":
            doc = DocxDocument()
            doc.add_heading("Generated Document", 0)
            for paragraph in content.split("\n\n"):
                doc.add_paragraph(paragraph)
            doc.save(file_path)
        else:
            doc = SimpleDocTemplate(file_path)
            styles = getSampleStyleSheet()
            flowables = [Paragraph(p, styles["Normal"]) for p in content.split("\n\n")]
            doc.build(flowables)

    elif file_format.lower() == "xlsx":
        rows = []
        if agent:
            agent_prompt = f"""
            Query the available data sources (preferably get_mysql_tool, then SQL database or document index) to extract structured data for this query:
            '{query}'

            ⚠️ VERY IMPORTANT:
            - Return ONLY structured data (plain CSV or JSON list of dicts)
            - First row = column headers
            - Do NOT include explanations or markdown
            - If no data is found, return exactly: Column1,Column2\nNo data available,,
            """
            result = agent.invoke({"input": agent_prompt})
            raw_output = result.get("output", "").strip()

            try:
                data_list = json.loads(raw_output)
                if isinstance(data_list, list) and len(data_list) > 0:
                    headers = list(data_list[0].keys())
                    rows.append(headers)
                    for row in data_list:
                        rows.append([row.get(h, "") for h in headers])
                else:
                    rows = [["No data available"]]
            except json.JSONDecodeError:
                lines = [line for line in raw_output.split("\n") if line.strip()]
                for line in lines:
                    rows.append([cell.strip() for cell in line.split(",")])

        else:
            rows = [["No data available"]]

        wb = Workbook()
        ws = wb.active
        for r_idx, row in enumerate(rows, start=1):
            for c_idx, value in enumerate(row, start=1):
                ws.cell(row=r_idx, column=c_idx, value=value)
        wb.save(file_path)

    else:
        raise ValueError("Unsupported format. Choose from 'pdf', 'docx', 'xlsx'.")

    return file_path

def determine_file_generation(query: str, conversation_history: str) -> tuple[bool, str | None, bool]:
    prompt = f"""
    You are an AI assistant analyzing a user's query to determine if it requires generating a file (e.g., a report, document, or spreadsheet) or a chart (e.g., bar, line, pie). Consider the query and conversation history to understand the user's intent. The query may not explicitly use words like 'generate', 'export', or 'chart' but may imply a need for a file or chart, such as requesting a report, table, or visual representation. If the query is ambiguous, assume no file or chart generation is needed unless clearly implied by context.

    Query: {query}
    Conversation History: {conversation_history}

    Respond with a valid JSON object (no extra text or code block markers like ```json):
    - "generate_file": boolean indicating if a file should be generated (true/false)
    - "file_format": string ("pdf", "docx", "xlsx") or null if no file is needed
    - "generate_chart": boolean indicating if a chart should be generated (true/false)

    Examples:
    - Query: "Give me a report of sales data" -> {{"generate_file": true, "file_format": "xlsx", "generate_chart": false}}
    - Query: "Create a bar chart of sales data" -> {{"generate_file": false, "file_format": null, "generate_chart": true}}
    - Query: "What is the return policy?" -> {{"generate_file": false, "file_format": null, "generate_chart": false}}
    - Query: "Summarize financials in a document with a chart" -> {{"generate_file": true, "file_format": "docx", "generate_chart": true}}
    - Query: "How many customers are there?" -> {{"generate_file": false, "file_format": null, "generate_chart": false}}
    """
    try:
        response = llm.invoke(prompt)
        result = response.content.strip()
        result = re.sub(r'^```json\s*|\s*```$', '', result, flags=re.MULTILINE)
        result = result.strip()
        parsed_result = json.loads(result)
        
        generate_file = parsed_result.get("generate_file", False)
        file_format = parsed_result.get("file_format", None)
        generate_chart = parsed_result.get("generate_chart", False)
        
        if generate_file and file_format not in ["pdf", "docx", "xlsx"]:
            print(f"Invalid file format '{file_format}' in LLM response. Defaulting to no file generation.")
            return False, None, generate_chart
        return generate_file, file_format, generate_chart

    except json.JSONDecodeError as e:
        print(f"JSON parsing error: {e}. Raw response: {result}. Defaulting to no file or chart generation.")
        return False, None, False
    except Exception as e:
        print(f"Error in determine_file_generation: {e}. Defaulting to no file or chart generation.")
        return False, None, False

def cleanup_generated_files(directory="."):
    try:
        for file in os.listdir(directory):
            if file.startswith("generated_report_") and file.endswith((".xlsx", ".pdf", ".docx")):
                file_path = os.path.join(directory, file)
                os.remove(file_path)
                print(f"Deleted file: {file_path}")
    except Exception as e:
        print(f"Error during cleanup of generated files: {e}")

def load_excel_to_db(excel_path, db_uri="sqlite:///structured_data.db"):
    table_name = os.path.splitext(os.path.basename(excel_path))[0]
    import pandas as pd
    from sqlalchemy import create_engine

    df = pd.read_excel(excel_path)
    df = df.fillna("")

    engine = create_engine(db_uri)
    df.to_sql(table_name, engine, if_exists='replace', index=False)
    
    return df, table_name

def build_full_user_history(username):
    if not username:
        return []
    return load_all_conversations(username)

def initialize_session_state():
    defaults = {
        "history": [],
        "generated_files": [],
        "generated_charts": [],
        "conversations": {},
        "active_conv": {},
        "authenticated": None,
        "user": None,
        "uploaded_files": [],
        "session_id": str(uuid.uuid4()),
        "memory": None
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

    if st.session_state.get("authenticated") and st.session_state.get("user"):
        if st.session_state.get("memory") is None:
            st.session_state.memory = ConversationBufferMemory(
                return_messages=True,
                memory_key="chat_history"
            )
            username = st.session_state.user
            all_history = build_full_user_history(username)
            for msg in all_history:
                if msg["role"] == "user":
                    st.session_state.memory.chat_memory.add_user_message(msg["content"])
                elif msg["role"] == "assistant":
                    st.session_state.memory.chat_memory.add_ai_message(msg["content"])
